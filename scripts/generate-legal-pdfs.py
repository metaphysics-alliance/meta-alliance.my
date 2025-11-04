from __future__ import annotations

import re
import subprocess
from pathlib import Path
from typing import Iterable

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parents[1]
PUBLIC_DIR = ROOT / "public"
LEGAL_DIR = PUBLIC_DIR / "legal"

MARKDOWN_SOURCES = [
    {
        "source": ROOT / "terms-of-service.md",
        "pdf": LEGAL_DIR / "terms-of-service.pdf",
        "docx": LEGAL_DIR / "terms-of-service.docx",
        "font": "Helvetica",
        "font_path": None,
    },
    {
        "source": ROOT / "privacy-policy.md",
        "pdf": LEGAL_DIR / "privacy-policy.pdf",
        "docx": LEGAL_DIR / "privacy-policy.docx",
        "font": "Helvetica",
        "font_path": None,
    },
    {
        "source": ROOT / "cookies-policy.md",
        "pdf": LEGAL_DIR / "cookies-policy.pdf",
        "docx": LEGAL_DIR / "cookies-policy.docx",
        "font": "Helvetica",
        "font_path": None,
    },
    {
        "source": ROOT / "terms-of-service-cn.md",
        "pdf": LEGAL_DIR / "terms-of-service-cn.pdf",
        "docx": LEGAL_DIR / "terms-of-service-cn.docx",
        "font": "SimHei",
        "font_path": Path("C:/Windows/Fonts/simhei.ttf"),
    },
    {
        "source": ROOT / "privacy-policy-cn.md",
        "pdf": LEGAL_DIR / "privacy-policy-cn.pdf",
        "docx": LEGAL_DIR / "privacy-policy-cn.docx",
        "font": "SimHei",
        "font_path": Path("C:/Windows/Fonts/simhei.ttf"),
    },
    {
        "source": ROOT / "cookies-policy-cn.md",
        "pdf": LEGAL_DIR / "cookies-policy-cn.pdf",
        "docx": LEGAL_DIR / "cookies-policy-cn.docx",
        "font": "SimHei",
        "font_path": Path("C:/Windows/Fonts/simhei.ttf"),
    },
]


def regenerate_markdown() -> None:
    script_path = ROOT / "scripts" / "write-legal-markdown.mjs"
    result = subprocess.run(
        ["node", str(script_path)],
        capture_output=True,
        text=True,
        check=True,
    )
    if result.stdout:
        print(result.stdout.strip())
    if result.stderr:
        print(result.stderr.strip())


def normalise_markdown(text: str) -> list[dict[str, str | int]]:
    """Convert Markdown-like content into structured blocks."""

    def clean(value: str) -> str:
        value = value.strip()
        if not value:
            return value
        value = re.sub(r"\*\*(.*?)\*\*", r"\\1", value)
        value = re.sub(r"\*(.*?)\*", r"\\1", value)
        value = value.replace("`", "")

        def _link_repl(match: re.Match[str]) -> str:
            label, href = match.group(1), match.group(2)
            if href.startswith("mailto:"):
                href = href.replace("mailto:", "")
            return f"{label} ({href})" if label != href else href

        value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", _link_repl, value)
        value = re.sub(r"\s+", " ", value)
        return value.strip()

    blocks: list[dict[str, str | int]] = []
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        striped = line.strip()
        indent_level = len(line) - len(line.lstrip(" "))
        if not striped:
            blocks.append({"type": "blank"})
            continue
        if striped in {"---", "***", "___"}:
            blocks.append({"type": "blank"})
            continue
        if line.lstrip().startswith(("-", "* ")):
            content = clean(line.lstrip()[1:].strip())
            blocks.append({"type": "bullet", "text": content, "indent": indent_level // 2})
            continue
        heading_match = re.match(r"^(#+)\s+(.*)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            content = clean(heading_match.group(2))
            blocks.append({"type": "heading", "level": level, "text": content})
            continue
        blocks.append({"type": "paragraph", "text": clean(line)})
    return blocks


def ensure_font(font_name: str, font_path: Path | None) -> None:
    if not font_path:
        return
    try:
        pdfmetrics.getFont(font_name)
        return
    except KeyError:
        pass
    if not font_path.exists():
        raise FileNotFoundError(f"Font path not found: {font_path}")
    pdfmetrics.registerFont(TTFont(font_name, str(font_path)))


def wrap_line(text: str, font_name: str, font_size: int, max_width: float) -> Iterable[str]:
    if not text:
        yield ""
        return
    has_space = bool(re.search(r"\s", text))
    units = text.split() if has_space else list(text)
    line = ""
    for unit in units:
        candidate = f"{line} {unit}".strip() if line else unit
        width = pdfmetrics.stringWidth(candidate, font_name, font_size)
        if width <= max_width:
            line = candidate
            continue
        if line:
            yield line
        line = unit
    if line:
        yield line


def render_pdf(blocks: list[dict[str, str | int]], output_path: Path, font_name: str) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    page_width, page_height = A4
    margin_x = 50
    margin_y = 60

    c = canvas.Canvas(str(output_path), pagesize=A4)

    def new_page() -> float:
        c.showPage()
        return page_height - margin_y

    y = page_height - margin_y
    max_width = page_width - 2 * margin_x

    for block in blocks:
        block_type = block.get("type")
        if block_type == "blank":
            y -= 12
            if y < margin_y:
                y = new_page()
            continue

        if block_type == "heading":
            level = int(block.get("level", 1))
            size = {1: 18, 2: 15, 3: 13}.get(level, 12)
            text = str(block.get("text", ""))
            c.setFont(font_name, size)
            for line in wrap_line(text, font_name, size, max_width):
                c.drawString(margin_x, y, line)
                y -= size * 1.5
                if y < margin_y:
                    y = new_page()
                    c.setFont(font_name, size)
            y -= 6
            continue

        if block_type == "bullet":
            size = 11
            c.setFont(font_name, size)
            indent_level = int(block.get("indent", 0))
            bullet_x = margin_x + indent_level * 18
            text_x = bullet_x + 12
            available_width = max_width - indent_level * 18 - 12
            lines = list(wrap_line(str(block.get("text", "")), font_name, size, available_width))
            if not lines:
                lines = [""]
            for idx, line in enumerate(lines):
                if idx == 0:
                    c.drawString(bullet_x, y, "•")
                    c.drawString(text_x, y, line)
                else:
                    c.drawString(text_x, y, line)
                y -= size * 1.5
                if y < margin_y:
                    y = new_page()
                    c.setFont(font_name, size)
            y -= 4
            continue

        if block_type == "paragraph":
            size = 11
            c.setFont(font_name, size)
            for line in wrap_line(str(block.get("text", "")), font_name, size, max_width):
                c.drawString(margin_x, y, line)
                y -= size * 1.5
                if y < margin_y:
                    y = new_page()
                    c.setFont(font_name, size)
            y -= 8
            continue

    c.save()


def render_docx(blocks: list[dict[str, str | int]], output_path: Path) -> None:
    doc = Document()
    bullet_styles = {0: 'List Bullet', 1: 'List Bullet 2', 2: 'List Bullet 3'}

    for block in blocks:
        block_type = block.get('type')
        if block_type == 'blank':
            doc.add_paragraph('')
            continue

        if block_type == 'heading':
            level = int(block.get('level', 1))
            text = str(block.get('text', ''))
            doc.add_heading(text, level=min(level, 4))
            continue

        if block_type == 'bullet':
            text = str(block.get('text', ''))
            indent_level = min(int(block.get('indent', 0)), 2)
            style = bullet_styles.get(indent_level, 'List Bullet 3')
            doc.add_paragraph(text, style=style)
            continue

        if block_type == 'paragraph':
            text = str(block.get('text', ''))
            doc.add_paragraph(text)
            continue

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    regenerate_markdown()
    for config in MARKDOWN_SOURCES:
        source = config["source"]
        pdf_target = config["pdf"]
        docx_target = config["docx"]
        font_name = config["font"]
        font_path = config.get("font_path")

        if font_path:
            ensure_font(font_name, font_path)

        blocks = normalise_markdown(source.read_text(encoding="utf-8"))
        render_pdf(blocks, pdf_target, font_name)
        render_docx(blocks, docx_target)
        print(f"[legal-docs] Generated {pdf_target.relative_to(ROOT)} and {docx_target.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
